"""Tests for the file-analysis pipeline (docparse + docanalysis + tools).

Network-free and dep-light: txt/csv need no extra packages; xlsx/pdf/docx tests
skip themselves if the matching optional lib isn't installed. `analyze()` is
driven by a fake provider so it never hits a real LLM.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from tradeflow.docanalysis import _extract_json, analyze  # noqa: E402
from tradeflow.docparse import ParseError, ParsedDoc, parse_file  # noqa: E402
from tradeflow.llm.base import LLMResponse  # noqa: E402
from tradeflow.tools.docanalysis import analyze_document, parse_document  # noqa: E402


def _tmp(suffix: str, *, text: str = None, bytes_: bytes = None) -> str:
    fd, path = tempfile.mkstemp(suffix=suffix)
    if bytes_ is not None:
        with os.fdopen(fd, "wb") as f:
            f.write(bytes_)
    else:
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(text or "")
    return path


class TestParse(unittest.TestCase):
    def test_txt_document(self):
        p = _tmp(".txt", text="合同存在风险条款\n第二行")
        d = parse_file(p)
        self.assertEqual(d.kind, "document")
        self.assertIn("风险", d.text)
        self.assertEqual(d.meta["format"], "text")
        self.assertEqual(d.meta["chars"], len("合同存在风险条款\n第二行"))

    def test_csv_table(self):
        p = _tmp(".csv", text="ASIN,Spend\nB001,10.5\nB002,20")
        d = parse_file(p)
        self.assertEqual(d.kind, "table")
        self.assertEqual(d.meta["rows_total"], 2)
        self.assertEqual(d.meta["columns"], ["ASIN", "Spend"])
        self.assertIn("ASIN", d.text)
        self.assertIn("B002", d.text)

    def test_unsupported_format(self):
        p = _tmp(".zip", text="x")
        with self.assertRaises(ParseError):
            parse_file(p)

    def test_xlsx_if_openpyxl(self):
        try:
            import openpyxl
        except ImportError:
            self.skipTest("openpyxl not installed")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["ASIN", "Spend"])
        ws.append(["B001", 10.5])
        ws.append(["B002", 20])
        p = _tmp(".xlsx")
        wb.save(p)
        d = parse_file(p)
        self.assertEqual(d.kind, "table")
        self.assertEqual(d.meta["rows_total"], 2)
        self.assertIn("B001", d.text)

    def test_docx_if_python_docx(self):
        try:
            import docx
        except ImportError:
            self.skipTest("python-docx not installed")
        doc = docx.Document()
        doc.add_paragraph("第一条：存在异常条款风险")
        doc.add_paragraph("第二条：正常条款")
        p = _tmp(".docx")
        doc.save(p)
        d = parse_file(p)
        self.assertEqual(d.kind, "document")
        self.assertIn("异常条款", d.text)


class TestExtractJson(unittest.TestCase):
    def test_plain(self):
        self.assertEqual(_extract_json('{"a": 1}'), {"a": 1})

    def test_fenced(self):
        self.assertEqual(_extract_json('```json\n{"a": 1}\n```'), {"a": 1})

    def test_prose_wrapped(self):
        self.assertEqual(
            _extract_json('结果如下：\n{"risks": [{"risk": "x"}]}\n以上。'),
            {"risks": [{"risk": "x"}]},
        )

    def test_braces_inside_string(self):
        self.assertEqual(_extract_json('{"q": "a { b } c"}'), {"q": "a { b } c"})

    def test_no_json(self):
        self.assertIsNone(_extract_json("这里没有任何 JSON"))


class _FakeProvider:
    """Stand-in LLMProvider: returns a canned LLMResponse regardless of input."""

    def __init__(self, text: str):
        self._t = text

    def complete(self, messages, system=None, tools=None):
        return LLMResponse(text=self._t)


class TestAnalyze(unittest.TestCase):
    def test_document_risks(self):
        parsed = ParsedDoc(text="合同文本…", kind="document", meta={"format": "text"})
        prov = _FakeProvider(json.dumps({
            "risks": [{"risk": "异常条款", "severity": "高",
                       "quote": "合同文本", "reason": "对己方不利"}],
        }))
        out = analyze(parsed, provider=prov)
        self.assertEqual(out["kind"], "document")
        self.assertEqual(len(out["risks"]), 1)
        self.assertEqual(out["risks"][0]["severity"], "高")
        self.assertEqual(out["risks"][0]["quote"], "合同文本")

    def test_table_insights(self):
        parsed = ParsedDoc(text="| ASIN | Spend |\n|---|---|\n| B001 | 10 |",
                           kind="table", meta={"format": "csv"})
        prov = _FakeProvider(json.dumps({
            "summary": "共 1 行",
            "insights": [{"type": "异常", "finding": "花费异常",
                          "evidence": "Spend=10", "action": "排查"}],
            "metrics": {"total_spend": "10"},
        }))
        out = analyze(parsed, provider=prov)
        self.assertEqual(out["summary"], "共 1 行")
        self.assertEqual(len(out["insights"]), 1)
        self.assertEqual(out["metrics"]["total_spend"], "10")

    def test_unparseable_returns_error(self):
        parsed = ParsedDoc(text="x", kind="document", meta={})
        out = analyze(parsed, provider=_FakeProvider("完全不是 JSON 的回复"))
        self.assertIn("error", out)
        self.assertEqual(out["raw"], "完全不是 JSON 的回复")

    def test_truncates_long_text(self):
        from tradeflow import docanalysis as da
        long = "X" * (da._DOC_MAX_CHARS_DEFAULT + 1000)
        parsed = ParsedDoc(text=long, kind="document", meta={"format": "text"})
        prompt = da._build_prompt(parsed, "")
        self.assertLess(len(prompt), len(long))


class TestTools(unittest.TestCase):
    def test_parse_document_returns_text(self):
        p = _tmp(".txt", text="hello contract world")
        out = json.loads(parse_document.run({"file_path": p}))
        self.assertEqual(out["kind"], "document")
        self.assertIn("hello", out["text"])

    def test_parse_document_caps_huge_text(self):
        p = _tmp(".txt", text="A" * 20000)
        out = json.loads(parse_document.run({"file_path": p}))
        self.assertLessEqual(len(out["text"]), 9000)  # 8000 cap + truncation note

    def test_parse_document_bad_path(self):
        out = json.loads(parse_document.run({"file_path": "C:/no/such/file.xyz"}))
        self.assertIn("error", out)

    def test_analyze_document_end_to_end(self):
        p = _tmp(".txt", text="合同含异常条款")
        # Monkeypatch docanalysis.analyze so no real provider is needed.
        from tradeflow.tools import docanalysis as tmod
        orig = tmod._analyze
        tmod._analyze = lambda parsed, focus="": {"risks": [{"risk": "x"}], "kind": parsed.kind}
        try:
            out = json.loads(analyze_document.run({"file_path": p}))
        finally:
            tmod._analyze = orig
        self.assertEqual(out["risks"][0]["risk"], "x")


if __name__ == "__main__":
    unittest.main()
